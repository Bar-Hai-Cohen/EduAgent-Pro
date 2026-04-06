import streamlit as st
import os
from openai import OpenAI
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from dotenv import load_dotenv

# --- 1. טעינת הגדרות ---
load_dotenv()

# מנגנון שבודק אם יש Secrets (לענן) או משתמש ב-ENV (למקומי) - מונע את השגיאה שקיבלת
try:
    if "OPENAI_API_KEY" in st.secrets:
        API_KEY = st.secrets["OPENAI_API_KEY"]
    else:
        API_KEY = os.getenv("OPENAI_API_KEY")
except Exception:
    API_KEY = os.getenv("OPENAI_API_KEY")

if not API_KEY:
    st.error("🔑 שגיאה: מפתח ה-API חסר.")
    st.stop()

client = OpenAI(api_key=API_KEY)


# --- 2. ניהול בסיס הנתונים (RAG) ---
@st.cache_resource(show_spinner="טוען את מסמכי המקור... 📚")
def init_vector_db():
    documents = []
    base_dir = "./data/EduAgent_Data"
    if not os.path.exists(base_dir): return None
    for root, dirs, files in os.walk(base_dir):
        for file in files:
            file_path = os.path.join(root, file)
            try:
                if file.endswith('.docx'):
                    doc = DocxDocument(file_path)
                    text = "\n".join([p.text for p in doc.paragraphs])
                elif file.endswith('.txt'):
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                else:
                    continue
                if text.strip():
                    documents.append(Document(page_content=text, metadata={"source": file}))
            except:
                continue
    if not documents: return None
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
    splits = text_splitter.split_documents(documents)
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")
    return FAISS.from_documents(splits, embeddings)


vector_db = init_vector_db()

# --- 3. אתחול זיכרון הצ'אט ---
if "messages" not in st.session_state:
    st.session_state.messages = []

# --- 4. ממשק המשתמש ---
st.set_page_config(page_title="EduAgent Chat", layout="wide", page_icon="💬")
st.title("🍎 EduAgent - צ'אט פדגוגי מבוסס מסמכים")

# תפריט צדדי
with st.sidebar:
    st.header("⚙️ אפשרויות")
    if st.button("נקה היסטוריית צ'אט"):
        st.session_state.messages = []
        st.rerun()
    st.info("הצ'אט זוכר את ההקשר של השיחה ומתבסס על 32 המסמכים שלך.")

# הצגת הודעות קודמות מהזיכרון
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# קלט מהמשתמש (צ'אט חי)
if prompt := st.chat_input("איך אוכל לעזור בתכנון הפדגוגי היום?"):

    # הצגת הודעת המשתמש
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    # שלב ה-RAG
    if not vector_db:
        st.error("בסיס הנתונים לא נטען.")
    else:
        with st.chat_message("assistant"):
            with st.status("🔍 סורק מסמכים רלוונטיים...") as status:
                # שליפה חכמה לפי השאלה האחרונה
                search_results = vector_db.similarity_search(prompt, k=10)
                context = "\n---\n".join([d.page_content for d in search_results])
                sources = list(set([d.metadata['source'] for d in search_results]))
                status.update(label=f"✅ נמצא מידע ב-{len(sources)} קבצים", state="complete")

            # בניית הפרומפט הכולל את ההיסטוריה
            history = "\n".join([f"{m['role']}: {m['content']}" for m in st.session_state.messages[-5:]])

            # בניית הפרומפט המקצועי לפי הגדרות EduAgent Pro
            full_system_prompt = f"""
                        אתה EduAgent Pro – מומחה בכיר בתחום הפדגוגיה, מומחה בתחום ההוראה ומנהל אסטרטגיית למידה. 
                        הנך בוט מומחה חינוך בהתאמה אישית. תפקידך לנתח את קלט המשתמש ולספק מענה פדגוגי מקיף הנשען אך ורק על המידע המאוזר ממאגר הנתונים, תוך ביצוע תפירת הקשר בין חמש שכבות מידע שונות להכנת פלט מותאם אישית למורה ולסביבתו. 
                        אתה מחולל מסמכי תכנון, הערכה ומערכים חינוכיים ברמת פירוט מקסימלית. עליך לייצר פלט ארוך מאוד (מינימום 3000 מילים במידת האפשר).

                        --- בסיס ידע מהמסמכים (Context) ---
                        {context}

                        --- כללי יסוד למניעת הזיות וחזרתיות ---
                        1. היצמדות מוחלטת למקורות (Anti-Hallucination): עליך לבסס כל המלצה, תוכנית פעולה או תפיסה אך ורק על המסמכים המקטולגים במאגר המידע. אין להמציא נהלים של משרד החינוך, מאפייני תלמידים, תיאוריות פדגוגיות או חזון בית-ספרי שאינם מופיעים בטקסט.
                        2. אינטגרציה ללא כפילות: הימנע מחזרה על אותם מסרים. נתח את הפרומפט וספק מענה קוהרנטי המשלב את המידע הקבוע עם המידע הספציפי של בית הספר והמורה לכדי תשובה רציפה ואופרטיבית.

                        --- מבנה פדגוגי קשיח - תהליך עיבוד וגיבוש המענה ---
                        על כל מענה שלך לעבור דרך פריזמה של חמש שכבות המידע ולהשתקף בהן:
                        - שכבה 1 - ביסוס תיאורטי ופדגוגי: עגן את תכנון ההוראה וההמלצות שלך בידע הרחב בחינוך הקיים במאגר. שלב תפיסות פדגוגיות, פילוסופיה של החינוך ומחקרים אקדמיים רלוונטיים (כגון למידה מבוססת מיומנויות, פדגוגיות הומניסטיות ואקטיביסטיות).
                        - שכבה 2 - עמידה בדרישות משרד החינוך: וודא כי הצעותיך תואמות במדויק את תוכניות הלימודים הרשמיות למקצוע הרלוונטי, חוזרי מנכ"ל ומסמכי ההנחיות למענה על שאלות בבחינות.
                        - שכבה 3 - פרקטיקה ויישום מעשי: שאב השראה מדוגמאות מעשיות של פרויקטים חינוכיים במאגר, הצע דרכי הוראה קונקרטיות, בניית תוכניות עבודה ומימוש של מיומנויות למידה בפועל.
                        - שכבה 4 - חיבור לזהות הבית-ספרית: התאם את המענה הפדגוגי לשפה הבית-ספרית הייחודית. שלב את חזון בית הספר, תפיסת עולמו הארגונית והמעשית, ואת הערכים והמפתחות האופרטיביים שהוגדרו כמרכזיים במוסד בו מלמד המורה.
                        - שכבה 5 - התאמה אישית: תפור את כלל המידע לכדי מענה ספציפי למורה הפונה. התחשב ב"אני מאמין" החינוכי שלו, במקצועות הלימוד שהוא מלמד, ובאפיון המדויק של הכיתות שלו (למשל: רמות למידה, אתגרים התנהגותיים, הרכב מגדרי או לקויות למידה).

                        --- מבנה הפלט הנדרש ---
                        הפלט חייב לכלול את התכנים הבאים:
                        1. המקצוע או המקצועות שעבורם נדרש הפלט.
                        2. נושא מדויק מתוך תוכנית הלימודים הרשמית של משרד החינוך.
                        3. התייחסות ספציפית לפרופיל הכיתה (מספר הכיתה, אפיון הכיתה מבחינה לימודית וחברתית, היום והשעה שבהם נלמד השיעור).
                        4. משך זמן השיעור וכל אחת מהפעילויות המוצעות בו.
                        5. חזון וערכים של בית הספר בהתאם למידע שבית הספר סיפק.
                        6. התאמה לתפיסת העולם והשפה המילולית של המורה בהתאם למידע שסיפק.
                        7. אופי מגוון של משימות שאפשר להשתמש בהם בפעילות החינוכית בהתאם לדוגמאות השונות שמסופקות במאגר המידע.

                        ענה בעברית רהוטה, מקצועית ומעמיקה מאוד.
                        """

            # שליחה ל-OpenAI עם הזיכרון
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": full_system_prompt},
                    *st.session_state.messages  # כל ההיסטוריה נשלחת ל-AI
                ],
                temperature=0.7
            )

            answer = response.choices[0].message.content
            st.markdown(answer)

            with st.expander("מקורות המידע לשאלה זו:"):
                st.write(", ".join(sources))

            # שמירת תשובת ה-AI לזיכרון
            st.session_state.messages.append({"role": "assistant", "content": answer})
